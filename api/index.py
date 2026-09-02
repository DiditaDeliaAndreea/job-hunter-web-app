import asyncio
from collections import Counter
import hashlib
import io
import json
import logging
import os
import re
import time
import uuid
from contextlib import asynccontextmanager
from datetime import date, datetime, timezone
import urllib.error
import urllib.parse
import urllib.request
from html import unescape
from typing import Any, Dict, List
from pathlib import Path

from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pypdf import PdfReader
from docx import Document
from google.api_core.exceptions import NotFound as GoogleNotFound
from openpyxl import Workbook

# Import CSV utilities
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from utils.csv_utils import append_jobs_to_csv, dismiss_expired_jobs_in_csv, dismiss_job_in_csv, export_jobs_to_csv, import_jobs_from_csv, import_local_jobs_from_csv, update_job_applied_in_csv, update_job_status_in_csv, update_job_url_in_csv
from utils import firebase_utils
from utils.hybrid_search import search_jobs as hybrid_search_jobs

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load local environment variables (.env)
load_dotenv()


def get_current_user(authorization: str | None = Header(default=None)) -> Dict[str, Any]:
    """Verify a Firebase ID token supplied by the browser."""
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Sign in is required")
    try:
        return firebase_utils.verify_id_token(authorization[7:].strip())
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail="Firebase server credentials are not configured") from exc
    except Exception as exc:
        logger.warning("Firebase token verification failed: %s", exc)
        raise HTTPException(status_code=401, detail="Invalid or expired sign-in session") from exc


def get_prompt_learning_context(user_id: str) -> str:
    preferences = firebase_utils.fetch_prompt_preferences(user_id)
    if not preferences:
        return "No saved tailoring preferences yet."
    return "\n".join(
        f"- Previous user preference ({record.get('interaction_type', 'tailor')}): {record.get('prompt', '')}"
        for record in preferences
        if str(record.get("prompt", "")).strip()
    )

@asynccontextmanager
async def validate_environment(_: FastAPI):
    """Validate required environment variables during application startup."""
    required_keys = ["GOOGLE_API_KEY"]
    missing = [key for key in required_keys if not os.getenv(key)]

    if missing:
        error_msg = f"Missing required environment variables: {', '.join(missing)}"
        logger.error(error_msg)
        raise RuntimeError(error_msg)

    logger.info("Environment variables validated")
    logger.info("Direct Google GenAI client ready")
    if os.getenv("OPENAI_API_KEY"):
        logger.info("OpenAI fallback configured")

    yield


app = FastAPI(lifespan=validate_environment)


@app.get("/")
async def root() -> JSONResponse:
    return JSONResponse({"service": "CareerMatch API", "status": "ok"})


@app.get("/health")
async def health() -> JSONResponse:
    return JSONResponse({"status": "ok"})

# Add CORS middleware to allow requests from frontend
frontend_urls = [
    origin.strip().rstrip("/")
    for origin in os.getenv("FRONTEND_URL", "http://localhost:3000").split(",")
    if origin.strip()
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=list(dict.fromkeys(frontend_urls + ["http://localhost:3000", "http://localhost:8000"])),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================================
# TARGET AND EXCLUDED ROLES CONFIGURATION
# ============================================================================

TARGET_ROLES = []

EXCLUDED_ROLES = []

TARGET_LOCATION = ""
ACTIVE_SEARCHES: Dict[str, Dict[str, Any]] = {}
CV_PROFILE_CACHE: Dict[str, str] = {}
CV_ANALYSIS_CONCURRENCY = 2
SEARCH_LOG_LIMIT = 25
SEARCH_RESULT_CACHE: Dict[tuple[str, str, str, bool], tuple[float, List[Dict[str, Any]]]] = {}
JOB_FIELDS = [
    "Job Title",
    "Company",
    "Location",
    "Posted Date",
    "Working Type",
    "Salary",
    "Fit Score (%)",
    "Match Reasons",
    "Missing Requirements",
    "Job Description",
    "Recommended CV",
    "CV Tailoring Recommendation",
    "Status",
    "URL",
    "Listing Source",
    "Official Listing Verified",
    "Official Listing URL",
    "Original Listing URL",
    "URL Check Status",
    "User Dismissed",
    "Applied",
    "User Status Override",
    "Verification Status",
    "Verification Notes",
    "Last Verified",
    "First Seen Date",
    "Is Actively Recruiting",
    "Extracted Skills",
    "Required Experience Years",
    "Required Experience Areas",
    "Seniority Level",
    "Must Have Requirements",
    "Nice To Have",
    "Embedding Text",
]


def _size_bucket(size: int) -> str:
    if size <= 1 * 1024 * 1024:
        return "0-1MB"
    if size <= 5 * 1024 * 1024:
        return "1-5MB"
    if size <= 10 * 1024 * 1024:
        return "5-10MB"
    return "over-10MB"


def _text_bucket(length: int) -> str:
    if length < 100:
        return "0-99"
    if length < 1000:
        return "100-999"
    if length < 5000:
        return "1000-4999"
    return "5000+"


def _file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    return suffix[1:] if suffix in {".pdf", ".docx"} else "other"


def _diagnostic_error_code(detail: object) -> str:
    message = str(detail).lower()
    if "scanned or image-only" in message or "no readable" in message or "selectable text" in message:
        return "no_extractable_text"
    if "corrupted" in message or "invalid" in message:
        return "invalid_file"
    if "rate" in message or "503" in message or "temporarily unavailable" in message:
        return "provider_unavailable"
    if "timeout" in message:
        return "timeout"
    return "internal_error"


def record_diagnostic_event(user_id: str, event: str, details: Dict[str, Any] | None = None) -> None:
    """Persist allowlisted operational data without user content."""
    if not firebase_utils.is_configured():
        return
    try:
        firebase_utils.save_diagnostic_log(user_id, event, details or {})
    except Exception as exc:  # Diagnostics must never break the user request.
        logger.warning("Could not persist diagnostic event %s: %s", event, type(exc).__name__)


def is_diagnostics_admin(user_id: str) -> bool:
    configured_admins = {
        value.strip()
        for value in os.getenv("ADMIN_UIDS", "").split(",")
        if value.strip()
    }
    return user_id in configured_admins
MIN_FULL_DESCRIPTION_LENGTH = 500


def is_expired_listing_text(text: str | None) -> bool:
    """Detect obsolete or closed listings that should be marked Expired."""
    if not text:
        return False
    normalized = text.lower()
    expired_markers = (
        "no longer accepting applications",
        "this job is no longer accepting applications",
        "applications are closed",
        "applications closed",
        "position is no longer open",
        "this role is no longer available",
        "job is no longer available",
        "oh no, this job is no longer available",
        "oh no this job is no longer available",
        "we've searched for similar jobs for you",
        "we have searched for similar jobs for you",
        "start a new search",
        "this position has been filled",
        "the position has been filled",
        "applications for this role have closed",
        "this listing is no longer active",
        "expired listing",
        "application deadline has passed",
    )
    return any(marker in normalized for marker in expired_markers)


def mark_stale_jobs_expired(jobs: List[Dict[str, Any]], stale_after_days: int = 21) -> bool:
    """Mark listings stale when their first-seen date has not been refreshed."""
    today = date.today()
    changed = False
    for job in jobs:
        if str(job.get("User Status Override") or "").strip().lower() == "yes":
            continue
        if str(job.get("Verification Status") or "").strip().casefold() == "expired":
            continue
        try:
            first_seen = date.fromisoformat(str(job.get("First Seen Date") or "").strip())
        except ValueError:
            job["First Seen Date"] = today.isoformat()
            changed = True
            continue
        if (today - first_seen).days > stale_after_days:
            job["Status"] = "Expired"
            job["Verification Status"] = "Expired"
            job["Verification Notes"] = f"Automatically marked stale after {stale_after_days} days without a source update."
            changed = True
    return changed


def normalize_job(job: Dict[str, Any]) -> Dict[str, Any] | None:
    """Keep only jobs with usable listing URLs and return the canonical CSV shape."""
    description = str(job.get("Job Description") or "").strip()
    if is_expired_listing_text(description):
        job["Status"] = "Expired"
        return None
    if (
        description.lower() in {"", "n/a", "not specified", "unknown"}
        or len(description) < MIN_FULL_DESCRIPTION_LENGTH
    ):
        return None

    source_url = str(job.get("Original Listing URL") or "").strip().strip("<>")
    official_url = str(job.get("Official Listing URL") or "").strip().strip("<>")
    submitted_url = str(job.get("URL") or "").strip().strip("<>")
    official_verified = str(job.get("Official Listing Verified") or "").strip().lower() == "yes"
    valid_url = lambda value: bool(re.match(r"^https?://[^\s]+$", value, re.IGNORECASE))

    if not valid_url(source_url):
        source_url = submitted_url if valid_url(submitted_url) else "Not specified"

    url = official_url if official_verified and valid_url(official_url) else source_url

    working_type = str(job.get("Working Type") or "Not specified").strip().lower()
    working_type_map = {
        "remote": "Remote",
        "hybrid": "Hybrid",
        "on-site": "On-site",
        "onsite": "On-site",
        "on site": "On-site",
    }

    def normalize_list(value: Any) -> str:
        if isinstance(value, list):
            return "; ".join(str(item).strip() for item in value if str(item).strip()) or "Not specified"
        return str(value or "Not specified").strip()

    extracted_skills = normalize_list(job.get("Extracted Skills") or job.get("extracted_skills"))
    required_experience_areas = normalize_list(
        job.get("Required Experience Areas") or job.get("required_experience_areas")
    )
    must_have_requirements = normalize_list(
        job.get("Must Have Requirements") or job.get("must_have_requirements")
    )
    nice_to_have = normalize_list(job.get("Nice To Have") or job.get("nice_to_have"))
    seniority_level = str(job.get("Seniority Level") or job.get("seniority_level") or "Not specified").strip()
    listing_status = str(job.get("Status") or "").strip().casefold()
    if listing_status in {"expired", "closed", "inactive", "not active"}:
        return None
    normalized_status = "Active" if listing_status in {"", "not specified", "open", "active"} else str(job.get("Status")).strip()
    first_seen_date = str(job.get("First Seen Date") or date.today().isoformat()).strip()
    embedding_text = " | ".join(
        value for value in (
            str(job.get("Job Title") or "").strip(),
            extracted_skills,
            required_experience_areas,
            seniority_level,
            must_have_requirements,
            nice_to_have,
            description,
        ) if value and value != "Not specified"
    )

    return {
        field: str(job.get(field) or "Not specified").strip()
        for field in JOB_FIELDS
    } | {
        "URL": url,
        "Job Description": description,
        "Original Listing URL": source_url,
        "URL Check Status": "Not checked",
        "Working Type": working_type_map.get(working_type, "Not specified"),
        "Status": normalized_status,
        "First Seen Date": first_seen_date,
        "Is Actively Recruiting": str(job.get("Is Actively Recruiting") or "Unknown").strip(),
        "Extracted Skills": extracted_skills,
        "Required Experience Years": str(
            job.get("Required Experience Years") or job.get("required_experience_years") or "Not specified"
        ).strip(),
        "Required Experience Areas": required_experience_areas,
        "Seniority Level": seniority_level,
        "Must Have Requirements": must_have_requirements,
        "Nice To Have": nice_to_have,
        "Embedding Text": embedding_text,
    }


def _aggregator_request(url: str, headers: Dict[str, str] | None = None) -> Dict[str, Any]:
    request = urllib.request.Request(url, headers=headers or {}, method="GET")
    with urllib.request.urlopen(request, timeout=20) as response:
        return json.loads(response.read().decode("utf-8"))


def _aggregator_post_request(url: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    body = json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(
        url, data=body, headers={"Content-type": "application/json"}, method="POST"
    )
    with urllib.request.urlopen(request, timeout=20) as response:
        return json.loads(response.read().decode("utf-8"))


def _posted_within_max_age(posted_date: str, max_posting_age_days: int) -> bool:
    """Enforce freshness using the aggregator's own timestamp rather than trusting AI-guessed dates."""
    if not posted_date or posted_date == "Not specified":
        return True
    try:
        parsed = datetime.fromisoformat(posted_date.replace("Z", "+00:00"))
    except ValueError:
        return True
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return (datetime.now(timezone.utc) - parsed).days <= max_posting_age_days


def _jsearch_date_posted(max_posting_age_days: int) -> str:
    if max_posting_age_days <= 1:
        return "today"
    if max_posting_age_days <= 3:
        return "3days"
    if max_posting_age_days <= 7:
        return "week"
    return "month"


def _jooble_posted_date(raw_job: Dict[str, Any]) -> str:
    """Normalize Jooble's "updated" timestamp to ISO 8601 for freshness filtering."""
    updated = str(raw_job.get("updated") or "").strip()
    if not updated:
        return "Not specified"
    try:
        return datetime.strptime(updated, "%Y-%m-%d %H:%M:%S").replace(tzinfo=timezone.utc).isoformat()
    except ValueError:
        return updated


def _map_aggregator_job(raw_job: Dict[str, Any], source: str) -> Dict[str, Any]:
    """Map JSearch/Jooble records into the existing normalized job contract."""
    if source == "jooble":
        title = raw_job.get("title") or "Not specified"
        company = raw_job.get("company") or "Not specified"
        location = raw_job.get("location") or "Not specified"
        url = raw_job.get("link") or "Not specified"
        description = raw_job.get("snippet") or "Not specified"
        posted_date = _jooble_posted_date(raw_job)
        salary = raw_job.get("salary") or "Not specified"
    else:
        title = raw_job.get("job_title") or "Not specified"
        company = raw_job.get("employer_name") or "Not specified"
        location = ", ".join(
            value for value in (raw_job.get("job_city"), raw_job.get("job_state"), raw_job.get("job_country")) if value
        ) or "Not specified"
        url = raw_job.get("job_apply_link") or raw_job.get("job_google_link") or "Not specified"
        description = raw_job.get("job_description") or "Not specified"
        posted_date = raw_job.get("job_posted_at_datetime_utc") or raw_job.get("job_posted_at") or "Not specified"
        salary = raw_job.get("job_min_salary") or raw_job.get("job_salary") or "Not specified"

    return {
        "Job Title": str(title),
        "Company": str(company),
        "Location": str(location),
        "Posted Date": str(posted_date),
        "Job Description": str(description),
        "Salary": str(salary),
        "URL": str(url),
        "Original Listing URL": str(url),
        "Listing Source": source.title(),
        "Official Listing Verified": "No",
        "Official Listing URL": "Not specified",
        "Status": "Active",
    }


async def fetch_aggregator_jobs(
    roles: List[str],
    location: str,
    max_posting_age_days: int,
) -> List[Dict[str, Any]]:
    """Fetch optional structured listings from JSearch and/or Jooble."""
    providers = {
        value.strip().lower()
        for value in os.getenv("JOB_AGGREGATORS", "").split(",")
        if value.strip()
    } & {"jsearch", "jooble"}
    if not providers:
        return []

    async def fetch_provider(provider: str) -> List[Dict[str, Any]]:
        try:
            if provider == "jooble":
                api_key = os.getenv("JOOBLE_API_KEY")
                if not api_key:
                    return []
                results = []
                for role in roles:
                    data = await asyncio.to_thread(
                        _aggregator_post_request,
                        f"https://jooble.org/api/{api_key}",
                        {"keywords": role, "location": location},
                    )
                    results.extend(_map_aggregator_job(job, provider) for job in data.get("jobs", []))
                # Jooble's free tier has a 500-request lifetime limit per key; keep this as a
                # supplementary source alongside JSearch rather than the sole provider.
                return [job for job in results if _posted_within_max_age(job["Posted Date"], max_posting_age_days)]

            rapid_key = os.getenv("RAPIDAPI_KEY")
            if not rapid_key:
                return []
            results = []
            for role in roles:
                params = urllib.parse.urlencode({
                    "query": role,
                    "num_pages": 1,
                    "country": "ie",
                    "date_posted": _jsearch_date_posted(max_posting_age_days),
                })
                data = await asyncio.to_thread(
                    _aggregator_request,
                    f"https://jsearch.p.rapidapi.com/search-v2?{params}",
                    {"X-RapidAPI-Key": rapid_key, "X-RapidAPI-Host": "jsearch.p.rapidapi.com"},
                )
                results.extend(_map_aggregator_job(job, provider) for job in data.get("data", []))
            return [job for job in results if _posted_within_max_age(job["Posted Date"], max_posting_age_days)]
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="ignore")[:300]
            logger.warning("%s job aggregation failed: HTTP %s %s - %s", provider, exc.code, exc.reason, body)
            return []
        except Exception as exc:
            logger.warning("%s job aggregation failed: %s", provider, type(exc).__name__)
            return []

    provider_results = await asyncio.gather(*(fetch_provider(provider) for provider in providers))
    return [job for result in provider_results for job in result]

def _url_tokens(value: str) -> set[str]:
    return {token for token in re.findall(r"[a-z0-9]+", value.lower()) if len(token) > 2}


def _is_linkedin_url(url: str) -> bool:
    return bool(re.search(r"(?:^|//)(?:[a-z]{2,3}\.)?linkedin\.com/", url, re.IGNORECASE))


def _url_is_expired_sync(url: str) -> bool:
    """Detect definitive HTTP removal or closure text for a saved listing URL."""
    if not url.startswith("http"):
        return False

    headers = {"User-Agent": "CareerMatch/1.0 job-listing-check"}
    try:
        request = urllib.request.Request(url, headers=headers, method="GET")
        with urllib.request.urlopen(request, timeout=5) as response:
            if response.headers.get_content_type() not in {"text/html", "application/xhtml+xml"}:
                return False
            page = response.read(512_000).decode("utf-8", errors="ignore")
    except urllib.error.HTTPError as exc:
        return exc.code in {404, 410}
    except (urllib.error.URLError, TimeoutError, ValueError):
        return False

    visible_text = re.sub(r"<script[^>]*>.*?</script>|<style[^>]*>.*?</style>|<[^>]+>", " ", page, flags=re.I | re.S)
    visible_text = unescape(re.sub(r"\s+", " ", visible_text))
    return is_expired_listing_text(visible_text)


def _check_url_sync(url: str, job_title: str, company: str, allow_client_rendered: bool = False) -> bool:
    """Check HTTP reachability and listing identity without blocking the event loop."""
    headers = {"User-Agent": "CareerMatch/1.0 job-listing-check"}
    try:
        request = urllib.request.Request(url, headers=headers, method="GET")
        with urllib.request.urlopen(request, timeout=5) as response:
            if not 200 <= response.status < 400:
                return False
            content_type = response.headers.get_content_type()
            if content_type not in {"text/html", "application/xhtml+xml"}:
                return True
            page = response.read(512_000).decode("utf-8", errors="ignore")
    except (urllib.error.HTTPError, urllib.error.URLError, TimeoutError, ValueError):
        return False

    visible_text = re.sub(r"<script[^>]*>.*?</script>|<style[^>]*>.*?</style>|<[^>]+>", " ", page, flags=re.I | re.S)
    visible_text = unescape(re.sub(r"\s+", " ", visible_text)).lower()
    title_tokens = _url_tokens(job_title)
    company_tokens = _url_tokens(company)
    title_matches = sum(token in visible_text for token in title_tokens)
    company_matches = sum(token in visible_text for token in company_tokens)
    soft_404_markers = (
        "job no longer available",
        "this job is no longer available",
        "page not found",
        "no jobs found",
        "job does not exist",
        "no longer accepting applications",
        "this job is no longer accepting applications",
        "applications are closed",
        "applications closed",
        "position is no longer open",
        "this role is no longer available",
        "job is no longer available",
        "oh no, this job is no longer available",
        "oh no this job is no longer available",
        "we've searched for similar jobs for you",
        "we have searched for similar jobs for you",
        "start a new search",
        "this position has been filled",
        "the position has been filled",
        "applications for this role have closed",
        "this listing is no longer active",
        "application deadline has passed",
    )
    url_text = unescape(url.lower().replace("-", " ").replace("_", " "))
    url_title_matches = sum(token in url_text for token in title_tokens)

    return (
        not any(marker in visible_text for marker in soft_404_markers)
        and (
            (
                title_matches >= max(1, min(3, len(title_tokens)))
                and company_matches >= 1
            )
            or (allow_client_rendered and url_title_matches >= max(1, min(3, len(title_tokens))))
        )
    )


async def validate_listing_job(job: Dict[str, Any]) -> Dict[str, Any] | None:
    """Prefer a reachable official URL but retain jobs with unavailable links."""
    source_url = job["Original Listing URL"]
    job_title = job.get("Job Title", "")
    company = job.get("Company", "")
    official_url = job.get("Official Listing URL", "Not specified")
    official_verified = str(job.get("Official Listing Verified", "")).lower() == "yes"

    source_expired, source_reachable, official_reachable = await asyncio.gather(
        asyncio.to_thread(_url_is_expired_sync, source_url)
        if source_url.startswith("http") else asyncio.sleep(0, result=False),
        asyncio.to_thread(_check_url_sync, source_url, job_title, company)
        if source_url.startswith("http") else asyncio.sleep(0, result=False),
        asyncio.to_thread(_check_url_sync, official_url, job_title, company, True)
        if official_verified and official_url != "Not specified"
        else asyncio.sleep(0, result=False),
    )

    if source_expired:
        logger.info(f"Dropping expired listing: {job_title} at {company} ({source_url})")
        return None

    if official_reachable:
        job["URL"] = official_url
        job["URL Check Status"] = "Official listing reachable"
    elif source_reachable:
        job["URL"] = source_url
        job["Official Listing Verified"] = "No"
        job["Official Listing URL"] = "Not specified"
        job["URL Check Status"] = "Source listing reachable"
    else:
        job["URL Check Status"] = "URL unavailable or blocked; job retained"
        logger.warning(f"Retaining job with unavailable URL: {job.get('Job Title')} ({source_url})")

    return job


async def validate_listing_urls(jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Validate all listing URLs concurrently and keep reachable jobs only."""
    checked_jobs = await asyncio.gather(*(validate_listing_job(job) for job in jobs))
    return [job for job in checked_jobs if job is not None]

def update_search_status(search_id: str, message: str, progress: int | None = None) -> None:
    """Track a running search so the frontend can show live progress updates."""
    status = ACTIVE_SEARCHES.setdefault(
        search_id,
        {"status": "running", "progress": 0, "message": message, "logs": []}
    )
    status["message"] = message
    if progress is not None:
        status["progress"] = progress
    status["logs"] = (status.get("logs") or []) + [message]
    if len(status["logs"]) > SEARCH_LOG_LIMIT:
        status["logs"] = status["logs"][-SEARCH_LOG_LIMIT:]
    logger.info(f"[search:{search_id}] {message}")


def parse_role_input(value: str | None, fallback: List[str]) -> List[str]:
    """Parse comma- or newline-separated roles and remove duplicate entries."""
    if not value or not value.strip():
        return fallback.copy()

    roles = [role.strip() for role in re.split(r"[,\n]", value) if role.strip()]
    return list(dict.fromkeys(roles)) or fallback.copy()


def _normalize_search_value(value: str) -> str:
    return re.sub(r"\s+", " ", value.strip())


ROLE_VARIANT_MAP = {
    "qa": ["QA Engineer", "Quality Assurance Engineer", "Software QA Engineer", "Software Test Engineer", "QA Analyst", "QA Automation Engineer", "Test Automation Engineer", "Quality Engineer"],
    "quality assurance": ["QA Engineer", "Quality Assurance Engineer", "Software QA Engineer", "QA Analyst", "Quality Engineer", "Software Test Engineer"],
    "support engineer": ["Technical Support Engineer", "Support Engineer", "Customer Support Engineer", "Application Support Engineer", "Technical Support Specialist", "Product Support Engineer", "Customer Experience Engineer"],
    "technical support": ["Technical Support Engineer", "Technical Support Specialist", "Application Support Engineer", "Customer Support Engineer", "Product Support Engineer"],
    "incident": ["Incident Response Analyst", "Incident Management Analyst", "Technical Support Engineer", "Operations Analyst", "Incident Analyst"],
    "operations analyst": ["Operations Analyst", "Operations Data Analyst", "Support Operations Analyst", "Technical Operations Analyst", "Business Operations Analyst"],
    "analyst": ["Analyst", "Operations Analyst", "Data Analyst", "Business Analyst", "Systems Analyst", "Technical Analyst", "Reporting Analyst", "Quality Analyst"],
    "customer success": ["Customer Success Engineer", "Technical Customer Success Specialist", "Implementation Consultant", "Implementation Specialist", "Customer Success Specialist"],
    "implementation": ["Implementation Consultant", "Implementation Specialist", "Technical Consultant", "Solutions Consultant"],
    "consultant": ["Technical Consultant", "Solutions Consultant", "Implementation Consultant", "Technical Solutions Consultant", "Customer Solutions Engineer"],
    "systems analyst": ["Systems Analyst", "Business Systems Analyst", "Technical Business Analyst", "Operations Analyst", "Product Operations Analyst"],
    "quality analyst": ["QA Analyst", "Quality Analyst", "Data Quality Analyst", "AI Quality Analyst", "Quality Assurance Analyst"],
    "test engineer": ["Software Test Engineer", "Test Automation Engineer", "QA Engineer", "Quality Assurance Engineer", "Software QA Engineer"],
    "ai support": ["AI Support Specialist", "AI Support Engineer", "AI Operations Analyst", "AI Operations Specialist", "Generative AI Support Specialist"],
    "product support": ["Product Support Engineer", "Product Support Specialist", "Customer Support Engineer", "Technical Support Engineer"],
    "technical manager": ["Technical Manager", "IT Manager", "Technology Manager", "Technical Operations Manager", "Technical Support Manager", "Systems Manager", "Service Delivery Manager", "Technical Project Manager"],
    "technology manager": ["Technology Manager", "IT Manager", "Technical Manager", "Technical Operations Manager", "Systems Manager", "Service Delivery Manager"],
}


def _expand_role_variants(role: str) -> List[str]:
    cleaned = _normalize_search_value(role)
    if not cleaned:
        return []
    variants = {cleaned}
    lower = cleaned.casefold()
    for key, values in ROLE_VARIANT_MAP.items():
        if key in lower or lower in key:
            variants.update(value for value in values if value.strip())
    replacements = [
        ("quality assurance engineer", "QA Engineer"),
        ("quality engineer", "QA Engineer"),
        ("software quality assurance engineer", "Software QA Engineer"),
        ("technical support specialist", "Technical Support Engineer"),
        ("support specialist", "Technical Support Engineer"),
        ("incident response specialist", "Incident Response Analyst"),
        ("incident management specialist", "Incident Management Analyst"),
    ]
    for source, replacement in replacements:
        if lower == source:
            variants.add(replacement)
    return [value for value in dict.fromkeys(variants) if value.strip()]


def extract_cv_role_suggestions(cv_summary: str) -> List[str]:
    """Read structured role suggestions from the CV analyzer response."""
    text = str(cv_summary or "").strip()
    try:
        parsed = json.loads(text)
    except (TypeError, json.JSONDecodeError):
        parsed = {}

    if isinstance(parsed, dict):
        role_candidates = parsed.get("matching_roles") or parsed.get("target_roles") or parsed.get("role_matches") or []
        if isinstance(role_candidates, list):
            roles = [
                _normalize_search_value(str(role))
                for role in role_candidates
                if str(role).strip()
            ]
            if roles:
                return roles

    matches = re.findall(r"(?:\b(?:QA|Quality Assurance|Technical Support|Customer Support|Application Support|Product Support|Systems Analyst|Operations Analyst|Incident Response|Incident Management|Customer Success|Implementation|Support|Software Quality|Test Automation|Quality Analyst|Data Analyst|Business Analyst)\b[^\n]{0,80})", text, flags=re.IGNORECASE)
    roles = []
    for match in matches:
        cleaned = _normalize_search_value(match)
        if cleaned and cleaned not in roles:
            roles.append(cleaned)
    return roles[:12]


def build_search_roles(
    target_roles: List[str],
    excluded_roles: List[str],
    cv_summaries: List[str],
) -> List[str]:
    """Combine explicit and CV-inferred roles, then remove excluded roles and include related role variants."""
    inferred_roles = []
    for cv_summary in cv_summaries:
        inferred_roles.extend(extract_cv_role_suggestions(cv_summary))

    expanded_roles: List[str] = []
    for role in [*target_roles, *inferred_roles]:
        expanded_roles.extend(_expand_role_variants(role))
    excluded = {role.casefold() for role in excluded_roles}
    filtered = [
        role
        for role in dict.fromkeys(expanded_roles)
        if role and role.casefold() not in excluded
    ]
    if filtered:
        return filtered
    return [role for role in dict.fromkeys([*target_roles, *inferred_roles]) if role and role.casefold() not in excluded]


def build_applied_job_preferences(jobs: List[Dict[str, Any]]) -> str:
    """Summarize the user's applied-job history for preference-aware search prompts."""
    applied_jobs = [
        job for job in jobs
        if str(job.get("Applied", "")).strip().lower() == "yes"
    ]
    if not applied_jobs:
        return "No applied-job history is available yet."

    role_categories = {
        "QA and testing": r"quality assurance|\bqa\b|test engineer|testing|tester|software quality|quality specialist|quality engineer",
        "Data and analytics": r"data analyst|data engineer|data operations|data quality|analytics|model operations|modelops",
        "Technical support": r"technical support|support engineer|support technician|application support|customer support|service delivery",
        "Operations": r"operations|assurance|capacity|content operations",
        "Business and systems analysis": r"business analyst|systems analyst|business applications",
        "Engineering and development": r"engineer|developer|software|automation|prompt",
        "Product and consulting": r"consultant|product specialist|product management",
    }
    category_counts = Counter()
    for job in applied_jobs:
        title = str(job.get("Job Title", "")).lower()
        matched_category = next(
            (category for category, pattern in role_categories.items() if re.search(pattern, title, re.IGNORECASE)),
            "Other",
        )
        category_counts[matched_category] += 1

    category_summary = ", ".join(
        f"{category} ({count})"
        for category, count in category_counts.most_common()
    )
    selected_roles = "\n".join(
        f"- {job.get('Job Title', 'Not specified')} — {job.get('Company', 'Not specified')}"
        for job in applied_jobs
    )
    return (
        f"The user has applied to {len(applied_jobs)} saved jobs. Their strongest demonstrated preferences "
        f"by role family are: {category_summary}. Use these as a strong ranking signal when selecting "
        "new roles, while still checking the CV, target roles, exclusions, location, freshness, and listing status. "
        "Applied-job history:\n" + selected_roles
    )


def update_verification_rows(rows: List[Dict[str, Any]], user_id: str) -> None:
    """Merge verification results into the saved CSV by title and company."""
    jobs = import_jobs_from_csv("job_matches.csv", user_id)
    result_map = {
        (str(row.get("Job Title", "")).strip().lower(), str(row.get("Company", "").strip().lower())): row
        for row in rows
    }
    for job in jobs:
        key = (job.get("Job Title", "").strip().lower(), job.get("Company", "").strip().lower())
        result = result_map.get(key)
        if not result:
            continue

        verification_status = str(result.get("Verification Status") or "Not found").strip()
        verification_notes = str(result.get("Verification Notes") or "Not specified").strip()
        if is_expired_listing_text(verification_notes) or is_expired_listing_text(str(result.get("Official Listing URL") or "")):
            verification_status = "Expired"
        if is_expired_listing_text(str(result.get("Job Title") or "")) or is_expired_listing_text(str(result.get("Company") or "")):
            verification_status = "Expired"
        urls_to_check = {
            str(result.get("Official Listing URL") or "").strip(),
            str(job.get("Official Listing URL") or "").strip(),
            str(job.get("URL") or "").strip(),
            str(job.get("Original Listing URL") or "").strip(),
        }
        linkedin_urls = {url for url in urls_to_check if _is_linkedin_url(url)}
        urls_to_check = linkedin_urls or urls_to_check
        if verification_status != "Expired" and any(_url_is_expired_sync(url) for url in urls_to_check):
            verification_status = "Expired"
            verification_notes = (
                "LinkedIn indicates that this job is no longer available."
                if linkedin_urls
                else "Listing page indicates that this job is no longer available."
            )

        job["Verification Status"] = verification_status
        job["Verification Notes"] = verification_notes or "Not specified"
        job["Last Verified"] = date.today().isoformat()
        official_url = str(result.get("Official Listing URL") or "").strip()
        if official_url.startswith("http"):
            job["Official Listing URL"] = official_url
            job["Official Listing Verified"] = "Yes"
            job["URL"] = official_url
            job["Listing Source"] = "Official company website"
    from utils.csv_utils import export_jobs_to_csv
    export_jobs_to_csv(jobs, "job_matches.csv", user_id)


async def run_saved_job_verification(verification_id: str, user_id: str) -> None:
    """Check saved roles in AI batches and persist results after each batch."""
    try:
        jobs = import_jobs_from_csv("job_matches.csv", user_id)
        jobs = [
            job for job in jobs
            if job.get("User Dismissed", "").lower() != "yes"
            and job.get("Applied", "").strip().lower() != "yes"
            and job.get("Verification Status", "").strip().lower() != "active"
            and job.get("User Status Override", "").lower() != "yes"
        ]
        batches = [jobs[index:index + 5] for index in range(0, len(jobs), 5)]
        for batch_index, batch in enumerate(batches, 1):
            progress = int(batch_index / max(1, len(batches)) * 90)
            update_search_status(verification_id, f"Verifying saved jobs batch {batch_index}/{len(batches)}...", progress)
            instructions = """
You verify saved job records using Google Search. Follow this order for every job:
1. Check the exact LinkedIn job posting first, including the direct LinkedIn URL when one is present.
2. Read the LinkedIn page status before checking any other source.
3. Only if LinkedIn does not provide a decisive result, check the exact job-board posting, then the
    employer's official careers website for the same exact title and location.
If LinkedIn says "No longer accepting applications", "applications are closed", "position is no
longer open", or gives any similar closure message, immediately classify the job as Expired. This
Expired result must not be overridden by a job-board copy, an official homepage, or an older search
snippet that appears active. Quote the closure wording and identify LinkedIn in Verification Notes.
Never mark Active from a generic careers homepage. Return ONLY a JSON array with Job Title, Company,
Verification Status, Verification Notes, and Official Listing URL.
"""
            prompt = "Verify these exact saved jobs:\n" + json.dumps([
                {"Job Title": job.get("Job Title"), "Company": job.get("Company"), "Location": job.get("Location"), "URL": job.get("URL"), "Original Listing URL": job.get("Original Listing URL")}
                for job in batch
            ], ensure_ascii=False)
            response = await call_gemini_with_retry(instructions, prompt, use_google_search=True, max_retries=3, initial_delay=5)
            try:
                verification_results = extract_json_array(response)
            except json.JSONDecodeError:
                update_search_status(verification_id, f"Batch {batch_index}: Gemini response was not JSON; trying OpenAI recovery...", progress)
                fallback_response = await call_openai_fallback(instructions, prompt, use_google_search=True)
                if not fallback_response:
                    raise
                verification_results = extract_json_array(fallback_response)
            update_verification_rows(verification_results, user_id)
            update_search_status(verification_id, f"Verified batch {batch_index}/{len(batches)} and saved results to CSV.", progress)
        ACTIVE_SEARCHES[verification_id]["status"] = "complete"
        ACTIVE_SEARCHES[verification_id]["progress"] = 100
        ACTIVE_SEARCHES[verification_id]["message"] = "Saved job verification complete."
    except Exception as exc:
        ACTIVE_SEARCHES[verification_id]["status"] = "error"
        ACTIVE_SEARCHES[verification_id]["error"] = str(exc)
        update_search_status(verification_id, f"Verification failed: {exc}", 0)

def extract_json_array(response: str) -> List[Dict[str, Any]]:
    """Extract the first valid JSON array from a model response."""
    cleaned = response.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.IGNORECASE | re.DOTALL).strip()

    decoder = json.JSONDecoder()
    for match in re.finditer(r"\[", cleaned):
        try:
            parsed, _ = decoder.raw_decode(cleaned[match.start():])
        except json.JSONDecodeError:
            continue
        if isinstance(parsed, list) and all(isinstance(item, dict) for item in parsed):
            return parsed

    raise json.JSONDecodeError("No valid JSON job array found", cleaned, 0)


async def run_search_pipeline(
    search_id: str,
    files: List[tuple[str, bytes]],
    target_roles: List[str],
    excluded_roles: List[str],
    reuse_cv_analysis: bool,
    user_id: str,
) -> None:
    """
    Main pipeline for analyzing CVs and finding jobs.

    This function now merges roles from the search form with the user's saved
    preferences from Firestore before starting the job search.
    """
    prefs = firebase_utils.fetch_role_preferences(user_id)
    saved_location = prefs.get("target_location", "")
    max_posting_age_days = int(prefs.get("max_posting_age_days", 7))
    try:
        total_files = len(files)
        analysis_semaphore = asyncio.Semaphore(CV_ANALYSIS_CONCURRENCY)

        async def process_cv(file_index: int, filename: str, content: bytes) -> tuple[str, str]:
            async with analysis_semaphore:
                content_hash = hashlib.sha256(content).hexdigest()
                update_search_status(
                    search_id,
                    f"Processing CV {file_index}/{total_files}: validating file...",
                    max(5, int((file_index - 1) / total_files * 35)),
                )
                validate_cv_filename(filename)
                cv_text = await extract_cv_text_from_bytes(filename, content)

                cv_summary = CV_PROFILE_CACHE.get(content_hash) if reuse_cv_analysis else None
                if cv_summary:
                    update_search_status(
                        search_id,
                        f"Reusing saved analysis for CV {file_index}/{total_files}...",
                        int(file_index / total_files * 35),
                    )
                else:
                    update_search_status(
                        search_id,
                        f"Analyzing CV {file_index}/{total_files} candidate profile...",
                        int(file_index / total_files * 35),
                    )
                    try:
                        cv_summary = await run_cv_analyzer_agent(cv_text)
                        CV_PROFILE_CACHE[content_hash] = cv_summary
                    except HTTPException as exc:
                        if exc.status_code != 503:
                            raise
                        # Never cache the fallback text as if it were real analysis, or every
                        # later search for this CV would keep reusing it and never retry.
                        cv_summary = (
                            "CV analysis was temporarily unavailable. Use the extracted CV content "
                            "directly for matching:\n" + cv_text[:12000]
                        )
                        update_search_status(
                            search_id,
                            f"CV {file_index}/{total_files} analysis unavailable; continuing with extracted CV text...",
                            int(file_index / total_files * 35),
                        )

                return filename, cv_summary

        analyzed_cvs = await asyncio.gather(
            *(process_cv(file_index, filename, content) for file_index, (filename, content) in enumerate(files, 1))
        )
        cv_summaries = [
            f"CV filename: {filename}\nCandidate profile:\n{cv_summary}"
            for filename, cv_summary in analyzed_cvs
        ]

        # Combine submitted roles with saved preferences for a comprehensive search
        final_target_roles = list(dict.fromkeys(target_roles + prefs.get("target_roles", [])))
        final_excluded_roles = list(dict.fromkeys(excluded_roles + prefs.get("excluded_roles", [])))

        expanded_target_roles = build_search_roles(final_target_roles, final_excluded_roles, cv_summaries)
        if expanded_target_roles:
            update_search_status(search_id, f"Expanded search to {len(expanded_target_roles)} related role signals based on the CV profile.", 55)
        combined_summary = "\n\n--- NEXT CV PROFILE ---\n\n".join(cv_summaries)
        update_search_status(search_id, "Searching live jobs in batches...", 55)
        matched_jobs = await run_incremental_job_finder(
            combined_summary,
            search_id=search_id,
            target_roles=expanded_target_roles,
            excluded_roles=final_excluded_roles,
            cv_names=[filename for filename, _ in files],
            user_id=user_id,
            target_location=saved_location,
            max_posting_age_days=max_posting_age_days,
        )

        if not matched_jobs:
            existing_jobs = import_jobs_from_csv("job_matches.csv", user_id)
            if existing_jobs:
                matched_jobs = existing_jobs
                update_search_status(search_id, f"Keeping {len(existing_jobs)} previous valid matches...", 72)
            else:
                matched_jobs = []
                update_search_status(search_id, "No matching jobs were found in this search.", 72)

        if matched_jobs:
            update_search_status(search_id, f"Persisting {len(matched_jobs)} matches to CSV...", 82)
            append_jobs_to_csv(matched_jobs, "job_matches.csv", user_id=user_id)
        else:
            update_search_status(search_id, "Skipping CSV write because no valid jobs were returned.", 82)

        update_search_status(search_id, "Exporting workbook to Excel...", 90)
        output = io.BytesIO()
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = 'Job Matches'
        headers = list(dict.fromkeys(field for job in matched_jobs for field in job.keys()))
        if headers:
            worksheet.append(headers)
            for job in matched_jobs:
                worksheet.append([job.get(header, '') for header in headers])
        workbook.save(output)
        output.seek(0)

        ACTIVE_SEARCHES[search_id]["status"] = "complete"
        ACTIVE_SEARCHES[search_id]["progress"] = 100
        ACTIVE_SEARCHES[search_id]["message"] = "Search complete. Excel file ready."
        ACTIVE_SEARCHES[search_id]["excel_bytes"] = output.getvalue()
        ACTIVE_SEARCHES[search_id]["filename"] = "job_matches.xlsx"
        ACTIVE_SEARCHES[search_id]["logs"] = (ACTIVE_SEARCHES[search_id].get("logs") or []) + ["Search complete. Excel file ready."]
        if len(ACTIVE_SEARCHES[search_id]["logs"]) > SEARCH_LOG_LIMIT:
            ACTIVE_SEARCHES[search_id]["logs"] = ACTIVE_SEARCHES[search_id]["logs"][-SEARCH_LOG_LIMIT:]
        logger.info(f"[search:{search_id}] Excel workbook generated successfully")

    except HTTPException as exc:
        ACTIVE_SEARCHES[search_id]["status"] = "error"
        ACTIVE_SEARCHES[search_id]["progress"] = 0
        ACTIVE_SEARCHES[search_id]["message"] = exc.detail
        ACTIVE_SEARCHES[search_id]["error"] = exc.detail
        update_search_status(search_id, f"Search failed: {exc.detail}", 0)
    except Exception as exc:  # pragma: no cover - defensive fallback
        ACTIVE_SEARCHES[search_id]["status"] = "error"
        ACTIVE_SEARCHES[search_id]["progress"] = 0
        ACTIVE_SEARCHES[search_id]["message"] = str(exc)
        ACTIVE_SEARCHES[search_id]["error"] = str(exc)
        update_search_status(search_id, f"Unexpected error: {exc}", 0)


async def call_gemini_with_retry(
    instructions: str,
    prompt: str,
    use_google_search: bool = False,
    max_retries: int = 5,
    initial_delay: int = 8
) -> str:
    """
    Calls Gemini API with retry logic for temporary spikes.

    When live web-search is required, this uses the Google GenAI SDK directly so the
    `google_search` tool is actually enabled for the request. The previous CrewAI 
    path ignored the `use_google_search` flag and therefore produced no live job data.
    """
    delay = initial_delay
    openai_attempted = False
    openai_error = ""

    def run_gemini_request() -> Any:
        from google import genai

        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            raise HTTPException(status_code=401, detail="Missing GOOGLE_API_KEY")

        client = genai.Client(api_key=api_key)
        return client.models.generate_content(
            model="gemini-flash-latest",
            contents=f"{instructions}\n\n{prompt}",
            config={"tools": [{"google_search": {}}]} if use_google_search else None,
        )

    for attempt in range(1, max_retries + 1):
        try:
            logger.info(f"Calling Gemini (attempt {attempt}/{max_retries})...")

            if use_google_search:
                response = await asyncio.wait_for(
                    asyncio.to_thread(run_gemini_request),
                    timeout=120,
                )

                text = getattr(response, "text", None)
                if text:
                    return str(text)

                candidates = getattr(response, "candidates", None) or []
                for candidate in candidates:
                    content = getattr(candidate, "content", None)
                    if not content:
                        continue
                    parts = getattr(content, "parts", []) or []
                    for part in parts:
                        part_text = getattr(part, "text", None)
                        if part_text:
                            return str(part_text)

                raise ValueError("Gemini returned an empty response for a Google Search call")

            response = await asyncio.wait_for(
                asyncio.to_thread(run_gemini_request),
                timeout=120,
            )
            text = getattr(response, "text", None)
            if text:
                return str(text)

            raise ValueError("Gemini returned an empty response for CV analysis")

        except Exception as e:
            err_str = str(e)
            if isinstance(e, asyncio.TimeoutError):
                err_str = "Google Search request timed out after 120 seconds"

            if attempt == 1 and os.getenv("OPENAI_API_KEY"):
                openai_attempted = True
                logger.warning("Gemini failed on the first attempt; OpenAI fallback started")
                openai_result = await call_openai_fallback(instructions, prompt, use_google_search)
                if openai_result is not None:
                    return openai_result
                openai_error = "OpenAI fallback failed or returned no usable response"
                raise HTTPException(
                    status_code=503,
                    detail=f"Gemini failed once; OpenAI fallback also failed: {openai_error}"
                )

            is_transient = any(
                code in err_str
                for code in ["503", "500", "429", "UNAVAILABLE", "RESOURCE_EXHAUSTED"]
            )

            if is_transient and attempt < max_retries:
                logger.warning(
                    f"⚠️ Server busy ({err_str[:60]}...). "
                    f"Retrying in {delay}s (Attempt {attempt}/{max_retries})..."
                )
                await asyncio.sleep(delay)
                delay *= 2
            else:
                logger.error(f"Gemini API call failed: {e}")
                fallback_detail = f"; {openai_error}" if openai_attempted else "; OpenAI fallback was not configured"
                raise HTTPException(
                    status_code=503,
                    detail=f"Gemini request failed after {attempt} attempt(s): {type(e).__name__}: {err_str[:240]}{fallback_detail}"
                )


async def call_openai_fallback(
    instructions: str,
    prompt: str,
    use_google_search: bool,
    model: str | None = None,
) -> str | None:
    """Call OpenAI and normalize web-search prose into the JSON contract when needed."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None

    def run_openai_request() -> Any:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        request: Dict[str, Any] = {
            "model": model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
            "input": f"{instructions}\n\n{prompt}",
        }
        if use_google_search:
            request["tools"] = [{"type": "web_search_preview"}]
        return client.responses.create(**request)

    try:
        logger.info("Calling OpenAI%s...", " with web search" if use_google_search else "")
        response = await asyncio.wait_for(
            asyncio.to_thread(run_openai_request),
            timeout=120,
        )
        text = getattr(response, "output_text", None)
        if text:
            logger.info("OpenAI fallback completed successfully")
            text = str(text)
            if not use_google_search:
                return text

            try:
                extract_json_array(text)
                return text
            except json.JSONDecodeError:
                logger.warning(
                    "OpenAI web-search response was not a JSON array; requesting JSON normalization"
                )

            def run_openai_normalization() -> Any:
                from openai import OpenAI

                client = OpenAI(api_key=api_key)
                normalization_instructions = (
                    "Convert the source response into the exact JSON array requested by the original "
                    "instructions. Preserve only factual job data from the source. Return ONLY a valid "
                    "JSON array of objects, with no markdown, explanation, or citations outside the JSON."
                )
                return client.responses.create(
                    model=model or os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
                    input=(
                        f"{normalization_instructions}\n\n"
                        f"ORIGINAL INSTRUCTIONS:\n{instructions}\n\n"
                        f"ORIGINAL REQUEST:\n{prompt}\n\n"
                        f"SOURCE RESPONSE:\n{text[:30000]}"
                    ),
                )

            normalized_response = await asyncio.wait_for(
                asyncio.to_thread(run_openai_normalization),
                timeout=120,
            )
            normalized_text = getattr(normalized_response, "output_text", None)
            if normalized_text:
                extract_json_array(str(normalized_text))
                logger.info("OpenAI web-search response normalized into JSON successfully")
                return str(normalized_text)
            raise ValueError("OpenAI JSON normalization returned an empty response")
        raise ValueError("OpenAI returned an empty response")
    except Exception as exc:
        logger.error(f"OpenAI fallback failed: {type(exc).__name__}: {exc}")
        return None


async def call_anthropic_extraction(instructions: str, prompt: str) -> str | None:
    """Use Claude for opt-in CV extraction without adding an SDK dependency."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return None

    def run_request() -> str:
        payload = json.dumps({
            "model": os.getenv("ANTHROPIC_EXTRACTION_MODEL", "claude-3-5-sonnet-latest"),
            "max_tokens": 3000,
            "temperature": 0,
            "system": instructions,
            "messages": [{"role": "user", "content": prompt}],
        }).encode("utf-8")
        request = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            },
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=120) as response:
            data = json.loads(response.read().decode("utf-8"))
        return "".join(
            str(block.get("text", ""))
            for block in data.get("content", [])
            if block.get("type") == "text"
        )

    try:
        return await asyncio.to_thread(run_request)
    except Exception as exc:
        logger.warning("Anthropic CV extraction failed: %s", type(exc).__name__)
        return None


async def call_cv_extraction_model(instructions: str, prompt: str) -> str:
    """Try the configured extractor first, then fall back to another provider."""
    preferred = os.getenv("CV_EXTRACTION_PROVIDER", "openai").strip().lower()
    providers = [preferred] + [provider for provider in ("openai", "anthropic", "gemini") if provider != preferred]
    for provider in providers:
        if provider == "openai":
            response = await call_openai_fallback(
                instructions,
                prompt,
                use_google_search=False,
                model=os.getenv("OPENAI_EXTRACTION_MODEL", "gpt-4o-mini"),
            )
        elif provider in {"anthropic", "claude"}:
            response = await call_anthropic_extraction(instructions, prompt)
        else:
            try:
                response = await call_gemini_with_retry(
                    instructions,
                    prompt,
                    use_google_search=False,
                    max_retries=2,
                    initial_delay=3,
                )
            except Exception as exc:
                logger.warning("Gemini CV extraction failed: %s", type(exc).__name__)
                response = None
        if response:
            logger.info("CV extraction completed with %s", provider)
            return response
    raise HTTPException(status_code=503, detail="All CV extraction providers failed or are unavailable")


async def call_job_match_model(instructions: str, prompt: str) -> str:
    """Score already-sourced listings against a candidate profile. Never used to discover jobs."""
    preferred = os.getenv("JOB_MATCH_PROVIDER", "openai").strip().lower()
    providers = [preferred] + [provider for provider in ("openai", "gemini") if provider != preferred]
    for provider in providers:
        if provider == "openai":
            response = await call_openai_fallback(
                instructions,
                prompt,
                use_google_search=False,
                model=os.getenv("OPENAI_MATCH_MODEL", "gpt-4o-mini"),
            )
        else:
            try:
                response = await call_gemini_with_retry(
                    instructions,
                    prompt,
                    use_google_search=False,
                    max_retries=3,
                    initial_delay=5,
                )
            except Exception as exc:
                logger.warning("Gemini job matching failed: %s", type(exc).__name__)
                response = None
        if response:
            logger.info("Job matching batch completed with %s", provider)
            return response
    raise HTTPException(status_code=503, detail="All job matching providers failed or are unavailable")


JOB_MATCH_FIT_RUBRIC = """
FIT SCORE RUBRIC — follow this exactly when setting "Fit Score (%)":
Start at 100 and apply the following deductions before assigning the score.

HARD REQUIREMENT PENALTIES (apply each that is unmet):
- Named proprietary tool/platform the candidate has no experience with (e.g. SAP, Salesforce,
  Hogan, Unibanks, Workday, ServiceNow, specific CMS): -20 per unmet tool, max -40
- Required years of experience the candidate clearly does not meet (e.g. "5+ years" when
  candidate has 1-2 years in that discipline): -20 per unmet requirement, max -30
- Mandatory degree or certification not present on the CV (e.g. "postgraduate required",
  "CPA required", "CISSP required"): -25
- Required domain-specific background the candidate has none of (e.g. "pharma GxP experience
  required", "financial services background required"): -20
- Automation or testing framework experience explicitly required but candidate's experience
  is clearly in a different form of automation (e.g. role requires Playwright/Selenium/
  Cypress test automation; candidate has workflow/process automation only): -20

OVERMATCHING GUARD — keyword presence alone is not a match:
- Do not award credit for a skill or tool if the CV only mentions it in passing or lists it
  without demonstrating professional depth. Only count skills with clear evidence of sustained
  use (multiple roles, quantified outcomes, or project ownership).
- Do not treat "QA" as equivalent to "test automation". Do not treat "Python scripting" as
  equivalent to "software engineering". Do not treat "data analysis" as equivalent to
  "BI/data analytics tooling experience". Match the depth, not just the keyword.

SCORING FLOOR:
- Any job with 2 or more unmet hard requirements must score 65% or below.
- Any job with a named proprietary tool requirement the candidate clearly lacks must score
  70% or below, regardless of other matches.

For "Match Reasons": write 2-4 sentences explaining the strongest genuine overlaps between
the candidate profile and this specific role. Be honest — do not pad with generic claims.

For "Missing Requirements": list every hard requirement from the job description that the
candidate does not clearly meet. Use a bullet list. Write "None identified" if there are
no meaningful gaps. This field must reflect the actual job description, not generic advice.
"""


async def run_job_matcher_agent(
    cv_summary: str,
    jobs: List[Dict[str, Any]],
    cv_names: List[str],
    applied_job_preferences: str,
) -> Dict[tuple[str, str], Dict[str, Any]]:
    """AGENT: Scores API-sourced listings against the candidate profile. Never discovers jobs itself."""
    if not jobs:
        return {}

    listings_block = "\n\n".join(
        f"[{index}] Job Title: {job.get('Job Title')}\n"
        f"Company: {job.get('Company')}\n"
        f"Location: {job.get('Location')}\n"
        f"Working Type: {job.get('Working Type')}\n"
        f"Description:\n{str(job.get('Job Description') or '')[:4000]}"
        for index, job in enumerate(jobs)
    )

    instructions = f"""
You are an expert recruiter scoring already-sourced job listings for one candidate. The listings
were retrieved from a structured job-search API, not by you — score only the numbered listings
provided below and never invent, add, or substitute a different job.

CANDIDATE PROFILE:
{cv_summary}

CANDIDATE CV FILES AVAILABLE (choose the best match, return the exact filename): [{", ".join(cv_names)}]

APPLIED-JOB PREFERENCE SIGNAL (use as a secondary hint, not a hard filter):
{applied_job_preferences}
{JOB_MATCH_FIT_RUBRIC}
Output ONLY a valid JSON array (no markdown, no ```json wrapper) with one object per listing index:
[
  {{
    "index": 0,
    "Fit Score (%)": "e.g., 85%",
    "Match Reasons": "2-4 sentences on genuine overlaps",
    "Missing Requirements": "Bullet list of unmet hard requirements, or None identified",
    "Extracted Skills": ["Hard skills, tools, frameworks, and certifications found in the listing"],
    "Required Experience Years": "Required years or Not specified",
    "Required Experience Areas": ["Core areas of required experience"],
    "Seniority Level": "Entry, Mid, Senior, Lead, Manager, or Not specified",
    "Must Have Requirements": ["Hard requirements that can disqualify a candidate"],
    "Nice To Have": ["Preferred but non-blocking qualifications"],
    "Embedding Text": "Concise normalized text combining title, skills, seniority, requirements, and responsibilities",
    "Recommended CV": "Exact filename from the candidate CV files list above",
    "CV Tailoring Recommendation": "Specific changes or emphasis for this job"
  }}
]
"""
    prompt = f"Score each numbered listing strictly against the candidate profile using the rubric above:\n\n{listings_block}"

    response = await call_job_match_model(instructions, prompt)
    try:
        scored = extract_json_array(response)
    except json.JSONDecodeError:
        logger.warning("Job matcher response was not parseable JSON; skipping scoring for this batch")
        return {}

    enrichment: Dict[tuple[str, str], Dict[str, Any]] = {}
    for item in scored:
        if not isinstance(item, dict):
            continue
        try:
            index = int(item.get("index"))
        except (TypeError, ValueError):
            continue
        if not 0 <= index < len(jobs):
            continue
        job = jobs[index]
        key = (str(job.get("Job Title", "")).strip().casefold(), str(job.get("Company", "")).strip().casefold())
        enrichment[key] = item
    return enrichment


async def run_cv_analyzer_agent(cv_text: str) -> str:
    """
    AGENT 1: Analyzes CV and builds a detailed candidate profile.
    
    Args:
        cv_text: Extracted text from the CV
        
    Returns:
        Candidate profile summary
    """
    logger.info("🤖 Agent 1 (CV Analyzer) starting - analyzing candidate profile...")
    
    location_context = f"in {TARGET_LOCATION}" if TARGET_LOCATION else "in their target market"
    instructions = f"""
    You are an expert technical recruiter matching candidates for roles in {location_context}.

    Analyze the provided resume and return strict JSON only.
    Required keys:
    - core_skills: array of technical skills and tools
    - seniority: string summarizing level and years of experience if stated
    - strengths: array of candidate specializations (e.g., QA testing, AI operations, triage)
    - matching_roles: array of likely job titles or role families that fit the candidate
    - languages: array of spoken or written languages explicitly mentioned in the CV (e.g., Romanian, English, Irish)
    - location: string for the candidate's current or target location if mentioned (e.g., "Dublin, Ireland")
    - years_of_experience: number when supported by explicit CV dates, otherwise null
    - primary_specializations: array of 2-5 evidence-based specializations
    - preferences: object with remote_or_hybrid, salary_expectation, and preferred_seniority keys; use null when absent
    - embedding_text: concise normalized text containing roles, skills, responsibilities, seniority, industries, and preferences

    Use evidence from the CV. Keep the output valid JSON with no markdown wrapper.
    """

    prompt = f"Analyze this CV and create a detailed candidate profile for {location_context} job matching:\n\n{cv_text}"
    
    response = await call_cv_extraction_model(instructions, prompt)
    logger.info("✅ Agent 1 completed: Candidate profile created")
    return response

@app.post("/api/jobs/recommend-cv")
async def recommend_cv_for_job(
    job_title: str = Form(...),
    company: str = Form(...),
    job_description: str = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Re-run CV matching for a saved job against the user's current CV library."""
    cvs = firebase_utils.fetch_cvs(current_user["uid"])
    if not cvs:
        raise HTTPException(status_code=400, detail="No CVs found. Upload at least one CV in My CVs first.")

    if len(cvs) == 1:
        recommended = cvs[0]["name"]
    else:
        cv_profiles: List[tuple[str, str]] = []
        for cv in cvs:
            try:
                record, content = firebase_utils.download_cv(current_user["uid"], cv["id"])
                text = await extract_cv_text_from_bytes(cv["name"], content)
                summary = await run_cv_analyzer_agent(text)
                cv_profiles.append((cv["name"], summary[:2000]))
            except Exception:
                continue

        if not cv_profiles:
            raise HTTPException(status_code=400, detail="Could not read CV files from storage.")

        profiles_block = "\n\n".join(
            f"=== CV: {name} ===\n{summary}" for name, summary in cv_profiles
        )
        instructions = """You are an expert recruiter. Given a job description and several candidate
        CV profiles, identify which CV is the strongest match for the role. Reply with ONLY the exact
        CV filename — no explanation, no punctuation, just the filename."""
        prompt = (
            f"Job Title: {job_title}\nCompany: {company}\n\n"
            f"Job Description:\n{job_description[:3000]}\n\n"
            f"Candidate CV Profiles:\n{profiles_block}\n\n"
            f"Which CV filename best matches this job?"
        )
        response = await call_gemini_with_retry(instructions, prompt)
        recommended = cvs[0]["name"]
        for cv in cvs:
            if cv["name"].strip().lower() in response.strip().lower():
                recommended = cv["name"]
                break

    all_jobs = firebase_utils.fetch_jobs(current_user["uid"])
    for job in all_jobs:
        if job.get("Job Title") == job_title and job.get("Company") == company:
            job["Recommended CV"] = recommended
    firebase_utils.replace_jobs(all_jobs, current_user["uid"])

    return JSONResponse({"recommended_cv": recommended})


@app.post("/api/jobs/tailor-cv")
async def tailor_cv_for_job(
    cv_file: UploadFile = File(...),
    job_title: str = Form(...),
    company: str = Form(...),
    job_description: str = Form(...),
    user_prompt: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Generate a truthful, keyword-aligned CV draft for one saved job."""
    content = await cv_file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The selected CV file is empty.")

    cv_text = await extract_cv_text_from_bytes(cv_file.filename or "uploaded-cv", content)
    if user_prompt.strip():
        firebase_utils.save_prompt_preference(
            current_user["uid"], user_prompt, job_title, company, "tailor"
        )
    instructions = """
You tailor a candidate's CV for one specific job. Preserve the candidate's actual experience and
never invent employers, dates, achievements, tools, certifications, responsibilities, or metrics.
Your primary goal is strong ATS alignment without keyword stuffing. Work in two explicit passes:

PASS 1 - JOB DESCRIPTION KEYWORD ANALYSIS:
Create an internal, prioritized keyword map from the job description. Identify meaningful ATS terms,
not filler words: required hard skills, software and tools, programming languages, methodologies,
certifications, domain terminology, job-specific responsibilities, seniority terms, deliverables,
and action verbs. Distinguish must-have keywords from useful secondary keywords and note important
phrases that an ATS may match exactly.

PASS 2 - CV EVIDENCE MATCHING AND REWRITE:
Audit the source CV against that keyword map. Include as many must-have and secondary keywords as
possible, but only when the source CV provides truthful evidence for them. Use exact job-description
wording when supported, and use accurate synonyms when they describe the same experience. Prioritize
the strongest supported matches near the top, in the summary, skills, and relevant experience
bullets. Rewrite existing bullets to lead with clear action verbs, scope, tools, and measurable
outcomes already present in the source CV. Never add a keyword merely because it would improve ATS
matching if the CV does not support it.

Return only a complete CV draft, ready to paste into the same document format as the source CV.
Preserve the source CV's exact section names, section order, heading wording, job-entry order,
bullet style, indentation, paragraph breaks, and spacing pattern wherever possible. Keep headings,
emphasis markers, capitalization, and date/location lines in the same style as the source. If the
source uses bold headings or labels, represent that emphasis with the same visible convention in
the output; do not replace it with a new heading style. Preserve the original bullet character
or marker rather than switching between bullets and paragraphs.

Keep the document ATS-safe and single-column: do not use tables, columns, text boxes, headers/footers,
icons, graphics, emojis, decorative symbols, unusual fonts, hyperlinks disguised as text, or layout
instructions. Do not rename sections to generic headings when the source already has clear headings.
Keep the wording concise and scannable. Do not include a cover letter, keyword list detached from
evidence, match score, editing commentary, or claims unsupported by the source CV. If a job
requirement is not supported, do not add it; leave it out or make the gap clear rather than guessing.
"""
    instructions += f"""
USER'S SAVED TAILORING PREFERENCES:
{get_prompt_learning_context(current_user['uid'])}

Treat these as style preferences gathered from the same user. Apply them only when they do not conflict
with the source CV, current job, or truthfulness requirements.
"""
    prompt = f"""
Target job: {job_title}
Company: {company}

USER'S DIRECT TAILORING REQUEST:
{user_prompt.strip() or 'No additional request. Follow the ATS and source-format instructions above.'}

JOB DESCRIPTION:
{job_description}

SOURCE CV:
{cv_text}

ATS CHECK BEFORE RETURNING:
- Every keyword added must be supported by the source CV.
- The most relevant supported job-description keywords must appear naturally in the summary,
  skills, or experience sections.
- The output must remain plain text, single-column, and machine-readable.
- Preserve the source CV's section names, ordering, bullets, line breaks, spacing, and emphasis style.
- Preserve factual names, dates, employers, and qualifications from the source CV.
"""
    try:
        tailored_cv = await call_gemini_with_retry(
            instructions,
            prompt,
            use_google_search=False,
            max_retries=3,
            initial_delay=5,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"CV tailoring failed: {type(exc).__name__}: {exc}")
        raise HTTPException(status_code=503, detail="Could not generate the tailored CV.") from exc

    return JSONResponse({"tailored_cv": tailored_cv})


@app.post("/api/ai/prompt-preferences")
async def save_ai_prompt_preference(
    prompt: str = Form(...),
    job_title: str = Form(""),
    company: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Enter a tailoring preference first.")
    record = firebase_utils.save_prompt_preference(
        current_user["uid"], prompt, job_title, company, "saved_preference"
    )
    return JSONResponse({"saved": True, "preference": record})


@app.post("/api/jobs/tailor-cv/chat")
async def ask_cv_tailoring_question(
    cv_file: UploadFile = File(...),
    job_title: str = Form(...),
    company: str = Form(...),
    job_description: str = Form(...),
    user_prompt: str = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Answer a direct CV-tailoring question for one job without generating a full CV."""
    if not user_prompt.strip():
        raise HTTPException(status_code=400, detail="Ask a tailoring question first.")

    content = await cv_file.read()
    if not content:
        raise HTTPException(status_code=400, detail="The selected CV file is empty.")

    cv_text = await extract_cv_text_from_bytes(cv_file.filename or "uploaded-cv", content)
    firebase_utils.save_prompt_preference(
        current_user["uid"], user_prompt, job_title, company, "question"
    )
    instructions = """
You are a CV tailoring advisor. Answer the user's direct question using the job description and source
CV. First identify meaningful ATS keywords and requirements in the job description, then check which
are supported by the CV. Give specific, actionable advice and exact suggested wording where useful.
Never invent experience, tools, employers, dates, certifications, metrics, or achievements. Preserve
the source CV's section names, bullet style, spacing pattern, and emphasis conventions. Be concise.
This is a one-request interaction: do not claim that the user's prompt trains, fine-tunes, or changes
the underlying model.
"""
    instructions += f"""
USER'S SAVED TAILORING PREFERENCES:
{get_prompt_learning_context(current_user['uid'])}

Use these as secondary style preferences, while prioritizing the current question and truthful CV evidence.
"""
    prompt = f"""
Target job: {job_title}
Company: {company}

USER QUESTION:
{user_prompt.strip()}

JOB DESCRIPTION:
{job_description}

SOURCE CV:
{cv_text}
"""
    try:
        answer = await call_gemini_with_retry(
            instructions,
            prompt,
            use_google_search=False,
            max_retries=3,
            initial_delay=5,
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"CV tailoring chat failed: {type(exc).__name__}: {exc}")
        raise HTTPException(status_code=503, detail="Could not answer the tailoring question.") from exc

    firebase_utils.save_prompt_preference(
        current_user["uid"], user_prompt, job_title, company, "question", answer
    )

    return JSONResponse({"answer": answer})


async def run_incremental_job_finder(
    cv_summary: str,
    search_id: str | None = None,
    target_roles: List[str] | None = None,
    excluded_roles: List[str] | None = None,
    cv_names: List[str] | None = None,
    user_id: str | None = None,
    target_location: str = "",
    max_posting_age_days: int = 7,
) -> List[Dict[str, Any]]:
    """
    AGENT 2: Fetches jobs from structured job-search APIs and scores them against the candidate.

    Job discovery is API-only (JSearch/Jooble via JOB_AGGREGATORS); AI is only used afterward to
    score, explain, and recommend a CV for the listings the APIs returned.

    Args:
        cv_summary: Candidate profile from Agent 1
        search_id: Optional active search ID for live progress updates
        target_location: User-configured location; falls back to CV-inferred location when empty.

    Returns:
        List of matched job opportunities
    """
    target_roles = target_roles or TARGET_ROLES
    excluded_roles = excluded_roles or EXCLUDED_ROLES
    cv_names = cv_names or ["Uploaded CV"]
    logger.info(f"🔎 Agent 2 (Job Finder) starting - querying structured job APIs for {len(target_roles)} target roles...")

    existing_jobs = import_jobs_from_csv("job_matches.csv", user_id)
    applied_job_preferences = build_applied_job_preferences(existing_jobs)
    saved_job_keys = {
        (
            str(job.get("Job Title", "")).strip().casefold(),
            str(job.get("Company", "")).strip().casefold(),
        )
        for job in existing_jobs
    }

    def keep_new_jobs(candidate_jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        unique_jobs = []
        for job in candidate_jobs:
            key = (
                str(job.get("Job Title", "")).strip().casefold(),
                str(job.get("Company", "")).strip().casefold(),
            )
            if key in saved_job_keys:
                continue
            saved_job_keys.add(key)
            unique_jobs.append(job)
        return unique_jobs

    def excluded_by_role(job: Dict[str, Any]) -> bool:
        title = str(job.get("Job Title", "")).casefold()
        return any(role.strip().casefold() in title for role in excluded_roles if role.strip())

    if search_id:
        update_search_status(search_id, "Querying structured job-search APIs...", 60)

    aggregator_location = target_location.strip() or TARGET_LOCATION or "Ireland"
    aggregator_jobs = await fetch_aggregator_jobs(target_roles, aggregator_location, max_posting_age_days)

    normalized_jobs = [normalize_job(job) for job in aggregator_jobs if isinstance(job, dict)]
    normalized_jobs = [job for job in normalized_jobs if job is not None]
    normalized_jobs = [job for job in normalized_jobs if not excluded_by_role(job)]
    normalized_jobs = await validate_listing_urls(normalized_jobs)
    normalized_jobs = keep_new_jobs(normalized_jobs)

    all_jobs: List[Dict[str, Any]] = []
    if not normalized_jobs:
        logger.info("✅ Agent 2 complete: No new jobs returned by the configured job-search APIs")
        if search_id:
            update_search_status(search_id, "No new jobs returned by the configured job-search APIs.", 95)
        return all_jobs

    if search_id:
        update_search_status(search_id, f"Scoring {len(normalized_jobs)} job(s) against your CV...", 75)

    match_batch_size = 10
    match_batches = [normalized_jobs[i:i + match_batch_size] for i in range(0, len(normalized_jobs), match_batch_size)]
    for batch_idx, batch in enumerate(match_batches, 1):
        try:
            enrichment = await run_job_matcher_agent(cv_summary, batch, cv_names, applied_job_preferences)
        except HTTPException as exc:
            logger.warning("Job matching failed for batch %s: %s", batch_idx, exc.detail)
            enrichment = {}

        for job in batch:
            key = (str(job.get("Job Title", "")).strip().casefold(), str(job.get("Company", "")).strip().casefold())
            match = enrichment.get(key)
            if not match:
                continue
            for field in (
                "Fit Score (%)",
                "Match Reasons",
                "Missing Requirements",
                "Required Experience Years",
                "Seniority Level",
                "Recommended CV",
                "CV Tailoring Recommendation",
            ):
                if match.get(field):
                    job[field] = str(match[field]).strip()
            for list_field in ("Extracted Skills", "Required Experience Areas", "Must Have Requirements", "Nice To Have"):
                value = match.get(list_field)
                if isinstance(value, list) and value:
                    job[list_field] = "; ".join(str(item).strip() for item in value if str(item).strip())
            embedding_text = match.get("Embedding Text")
            if embedding_text:
                job["Embedding Text"] = str(embedding_text).strip()

        if search_id:
            progress_pct = 75 + int((batch_idx / len(match_batches)) * 20)
            update_search_status(
                search_id,
                f"Scored batch {batch_idx}/{len(match_batches)} ({len(batch)} job(s)).",
                progress_pct,
            )

    all_jobs.extend(normalized_jobs)
    if user_id:
        append_jobs_to_csv(normalized_jobs, "job_matches.csv", user_id=user_id)
    else:
        append_jobs_to_csv(normalized_jobs, "job_matches.csv")

    logger.info(f"✅ Agent 2 complete: Found and scored {len(all_jobs)} job(s) via structured job-search APIs")
    if search_id:
        update_search_status(search_id, f"Job search complete: {len(all_jobs)} job(s) found via structured APIs.", 95)
    return all_jobs


def validate_cv_filename(filename: str | None) -> None:
    """Validate the filename of an uploaded CV file."""
    if not filename:
        raise HTTPException(status_code=400, detail="Missing filename")

    filename_lower = filename.lower()
    if not filename_lower.endswith(('.pdf', '.docx')):
        raise HTTPException(
            status_code=400,
            detail="Only PDF and DOCX files are supported"
        )

    logger.info(f"File {filename} passed validation checks")


def has_extractable_cv_text(text: str) -> bool:
    """Return whether extracted content contains enough readable CV text to analyze."""
    return len(re.findall(r"[A-Za-z0-9]", text)) >= 20


async def extract_cv_text_from_bytes(filename: str, content: bytes) -> str:
    """Extract text from a PDF or DOCX payload without needing the UploadFile object."""
    if not content:
        raise HTTPException(status_code=400, detail="Empty file uploaded")

    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        try:
            pdf = PdfReader(io.BytesIO(content))
            if not pdf.pages:
                raise HTTPException(status_code=422, detail="PDF has no pages")

            cv_text = ""
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    cv_text += text
                    logger.debug(f"Extracted {len(text)} chars from PDF page {page_num}")

            if not has_extractable_cv_text(cv_text):
                logger.warning("PDF extracted but contains no readable text")
                raise HTTPException(
                    status_code=422,
                    detail="PDF contains no readable text. It may be scanned or image-only; upload a text-based PDF or DOCX.",
                )

            logger.info(f"Successfully extracted {len(cv_text)} chars from PDF")
            return cv_text

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PDF parsing failed: {type(e).__name__}: {e}")
            raise HTTPException(status_code=422, detail="Invalid or corrupted PDF file")

    elif filename_lower.endswith('.docx'):
        try:
            doc = Document(io.BytesIO(content))
            if not doc.paragraphs:
                raise HTTPException(
                    status_code=422,
                    detail="DOCX contains no readable text. Check that the document contains selectable text.",
                )

            cv_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    cv_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            cv_text += cell.text + " "
                    cv_text += "\n"

            if not has_extractable_cv_text(cv_text):
                logger.warning("DOCX extracted but contains no readable text")
                raise HTTPException(
                    status_code=422,
                    detail="DOCX contains no readable text. Check that the document contains selectable text.",
                )

            logger.info(f"Successfully extracted {len(cv_text)} chars from DOCX")
            return cv_text

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX parsing failed: {type(e).__name__}: {e}")
            raise HTTPException(status_code=422, detail="Invalid or corrupted DOCX file")

    raise HTTPException(status_code=400, detail="Unsupported file type")


@app.post("/api/search/start")
async def start_search(
    files: List[UploadFile] = File(default=[]),
    file: UploadFile | None = File(default=None),
    target_roles: str | None = Form(default=None),
    excluded_roles: str | None = Form(default=None),
    reuse_cv_analysis: bool = Form(default=True),
    cv_ids: str | None = Form(default=None),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Start a background job search and return a live search ID for polling."""
    uploaded_files = list(files)
    if file is not None:
        uploaded_files.append(file)
    file_payloads = [(uploaded_file.filename or "unknown.pdf", await uploaded_file.read()) for uploaded_file in uploaded_files]
    selected_cv_ids = [cv_id.strip() for cv_id in (cv_ids or "").split(",") if cv_id.strip()]
    for cv_id in selected_cv_ids:
        try:
            cv_record, cv_content = firebase_utils.download_cv(current_user["uid"], cv_id)
        except KeyError as exc:
            raise HTTPException(status_code=404, detail=f"Selected CV not found: {cv_id}") from exc
        file_payloads.append((str(cv_record.get("name") or "uploaded-cv.pdf"), cv_content))
    if not file_payloads:
        raise HTTPException(status_code=400, detail="Upload at least one PDF or DOCX CV file")

    parsed_target_roles = parse_role_input(target_roles, TARGET_ROLES)
    parsed_excluded_roles = parse_role_input(excluded_roles, EXCLUDED_ROLES)
    search_id = f"search-{int(time.time() * 1000)}-{uuid.uuid4().hex[:8]}"
    ACTIVE_SEARCHES[search_id] = {
        "status": "queued",
        "progress": 0,
        "message": "Starting job search...",
        "logs": ["Starting job search..."],
        "total_files": len(file_payloads),
        "reuse_cv_analysis": reuse_cv_analysis,
        "user_id": current_user["uid"],
    }
    asyncio.create_task(run_search_pipeline(
        search_id,
        file_payloads,
        parsed_target_roles,
        parsed_excluded_roles,
        reuse_cv_analysis,
        current_user["uid"],
    ))
    return JSONResponse({
        "search_id": search_id,
        "status": "queued",
        "message": f"Search started for {len(file_payloads)} CV(s)",
    })


@app.get("/api/cvs")
async def get_cvs(current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    try:
        return JSONResponse({"cvs": firebase_utils.fetch_cvs(current_user["uid"])})
    except Exception as exc:
        logger.error("Error fetching CVs: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to fetch imported CVs") from exc


@app.post("/api/cvs")
async def upload_cvs(
    files: List[UploadFile] = File(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    saved_cvs = []
    for uploaded_file in files:
        filename = uploaded_file.filename or "uploaded-cv.pdf"
        if not filename.lower().endswith((".pdf", ".docx")):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX CVs are supported")
        content = await uploaded_file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Each CV must be 10MB or smaller")
        try:
            saved_cvs.append(firebase_utils.save_cv(
                current_user["uid"],
                filename,
                uploaded_file.content_type or "application/octet-stream",
                content,
            ))
        except GoogleNotFound as exc:
            raise HTTPException(
                status_code=503,
                detail="Firebase Storage is not enabled for this project. Open Firebase Console > Storage and click Get started, then retry.",
            ) from exc
    return JSONResponse({"cvs": saved_cvs})


@app.delete("/api/cvs/{cv_id}")
async def delete_cv(cv_id: str, current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    try:
        firebase_utils.delete_cv(current_user["uid"], cv_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="CV not found") from exc
    return JSONResponse({"deleted": True, "cv_id": cv_id})


@app.patch("/api/cvs/{cv_id}")
async def rename_cv(cv_id: str, name: str = Form(...), current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    new_name = name.strip()
    if not new_name.lower().endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="CV name must end with .pdf or .docx")
    if "/" in new_name or "\\" in new_name:
        raise HTTPException(status_code=400, detail="CV name cannot contain folder separators")
    try:
        record = firebase_utils.rename_cv(current_user["uid"], cv_id, new_name)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="CV not found") from exc
    return JSONResponse({"cv": record})


@app.get("/api/cvs/{cv_id}/content")
async def get_cv_content(cv_id: str, current_user: Dict[str, Any] = Depends(get_current_user)) -> StreamingResponse:
    try:
        record, content = firebase_utils.download_cv(current_user["uid"], cv_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="CV not found") from exc
    return StreamingResponse(
        io.BytesIO(content),
        media_type=str(record.get("content_type") or "application/octet-stream"),
        headers={"Content-Disposition": f"inline; filename=\"{record.get('name', 'cv')}\""},
    )


@app.get("/api/search/status/{search_id}")
async def get_search_status(search_id: str, current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    """Return the current progress and log stream for an active search."""
    status = ACTIVE_SEARCHES.get(search_id)
    if not status or status.get("user_id") != current_user["uid"]:
        raise HTTPException(status_code=404, detail="Search not found")

    return JSONResponse(
        {
            "search_id": search_id,
            "status": status.get("status", "running"),
            "progress": status.get("progress", 0),
            "message": status.get("message", "Processing..."),
            "logs": status.get("logs", []),
            "download_url": f"/api/search/result/{search_id}" if status.get("status") == "complete" else None,
            "error": status.get("error"),
        }
    )


@app.get("/api/search/result/{search_id}")
async def get_search_result(search_id: str, current_user: Dict[str, Any] = Depends(get_current_user)) -> StreamingResponse:
    """Download the Excel workbook for a completed search."""
    status = ACTIVE_SEARCHES.get(search_id)
    if not status or status.get("user_id") != current_user["uid"] or status.get("status") != "complete":
        raise HTTPException(status_code=404, detail="Search not complete")

    excel_bytes = status.get("excel_bytes")
    if not excel_bytes:
        raise HTTPException(status_code=404, detail="Excel file not available")

    return StreamingResponse(
        io.BytesIO(excel_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={status.get('filename', 'job_matches.xlsx')}"}
    )


@app.get("/api/jobs")
async def get_jobs(current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    """
    Fetch all jobs from CSV file.
    
    Returns:
        JSON array of job objects
    """
    try:
        jobs = import_jobs_from_csv("job_matches.csv", current_user["uid"])
        stale_after_days = max(14, int(os.getenv("STALE_JOB_DAYS", "21")))
        if mark_stale_jobs_expired(jobs, stale_after_days):
            export_jobs_to_csv(jobs, "job_matches.csv", user_id=current_user["uid"])
        jobs = [job for job in jobs if job.get("User Dismissed", "").lower() != "yes"]
        
        if not jobs:
            return JSONResponse(
                status_code=200,
                content={"jobs": [], "message": "No jobs found. Upload a CV to search."}
            )
        
        logger.info(f"✅ Fetched {len(jobs)} jobs from CSV")
        return JSONResponse(
            status_code=200,
            content={"jobs": jobs, "total": len(jobs)}
        )

    except Exception as e:
        logger.error(f"Error fetching jobs: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to fetch jobs from database"
        )


async def call_cohere_rerank(query: str, documents: List[str]) -> List[Dict[str, Any]] | None:
    """Use Cohere's dedicated reranker when explicitly configured."""
    api_key = os.getenv("COHERE_API_KEY")
    if not api_key or not documents:
        return None

    def run_request() -> List[Dict[str, Any]]:
        payload = json.dumps({
            "model": os.getenv("COHERE_RERANK_MODEL", "rerank-v3.5"),
            "query": query,
            "documents": documents,
            "top_n": len(documents),
            "return_documents": False,
        }).encode("utf-8")
        request = urllib.request.Request(
            "https://api.cohere.com/v2/rerank",
            data=payload,
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=60) as response:
            return json.loads(response.read().decode("utf-8")).get("results", [])

    try:
        return await asyncio.to_thread(run_request)
    except Exception as exc:
        logger.warning("Cohere reranking failed: %s", type(exc).__name__)
        return None


async def rerank_jobs(profile: str, jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Rerank a small pool with a configured provider, retaining results on failure."""
    candidates = jobs[:10]
    candidate_text = "\n\n".join(
        f"JOB {index}: {job.get('Job Title', 'Not specified')} at {job.get('Company', 'Not specified')}\n"
        f"Skills: {job.get('Extracted Skills', 'Not specified')}\n"
        f"Requirements: {job.get('Must Have Requirements', job.get('Missing Requirements', 'Not specified'))}\n"
        f"Description: {job.get('Job Description', 'Not specified')}"
        for index, job in enumerate(candidates)
    )
    instructions = """
You are a strict technical recruiting reranker. Compare the candidate profile to each job.
Return ONLY a JSON array with one object per candidate, using the candidate's JOB number:
[{"job_index": 0, "match_score": 0, "why_match": "", "skill_gaps": [""]}]
Use a 0-100 score. Cite genuine overlaps, identify concrete missing hard requirements, and never
infer experience that is not supported by the candidate profile.
"""
    prompt = f"CANDIDATE PROFILE:\n{profile[:12000]}\n\nRETRIEVED JOBS:\n{candidate_text}"
    provider = os.getenv("RERANK_PROVIDER", "gemini").strip().lower()
    try:
        if provider == "cohere":
            cohere_results = await call_cohere_rerank(profile, [
                f"{job.get('Job Title', '')} at {job.get('Company', '')}: {job.get('Job Description', '')}"
                for job in candidates
            ])
            if not cohere_results:
                return jobs
            ordered = []
            for result in sorted(cohere_results, key=lambda item: float(item.get("relevance_score", 0)), reverse=True):
                index = int(result.get("index", -1))
                if 0 <= index < len(candidates):
                    enriched = dict(candidates[index])
                    enriched["Rerank Score"] = f"{float(result.get('relevance_score', 0)):.4f}"
                    ordered.append(enriched)
            ordered.extend(job for index, job in enumerate(candidates) if not any(job is item for item in ordered))
            return ordered + jobs[10:]
        if provider == "openai":
            response = await call_openai_fallback(
                instructions,
                prompt,
                use_google_search=False,
                model=os.getenv("OPENAI_RERANK_MODEL", "gpt-4o"),
            )
        elif provider in {"anthropic", "claude"}:
            response = await call_anthropic_extraction(instructions, prompt)
        else:
            response = await call_gemini_with_retry(instructions, prompt, use_google_search=False, max_retries=2, initial_delay=2)
        if not response:
            return jobs
        rankings = extract_json_array(response)
    except Exception as exc:
        logger.warning("Job reranking unavailable; retaining hybrid retrieval order: %s", type(exc).__name__)
        return jobs

    ranking_map = {
        int(item["job_index"]): item
        for item in rankings
        if isinstance(item, dict) and str(item.get("job_index", "")).isdigit()
        and 0 <= int(item["job_index"]) < len(candidates)
    }
    reranked = []
    scored_indexes = set()
    for index, job in enumerate(candidates):
        ranking = ranking_map.get(index)
        if ranking:
            enriched_job = dict(job)
            raw_score = str(ranking.get("match_score", 0)).replace("%", "").strip()
            try:
                match_score = int(float(raw_score))
            except ValueError:
                match_score = 0
            enriched_job["Fit Score (%)"] = f"{max(0, min(100, match_score))}%"
            enriched_job["Match Reasons"] = str(ranking.get("why_match") or "Not specified")
            enriched_job["Missing Requirements"] = "; ".join(
                str(gap).strip() for gap in ranking.get("skill_gaps", []) if str(gap).strip()
            ) or "None identified"
            reranked.append((match_score, enriched_job))
            scored_indexes.add(index)
    reranked.sort(key=lambda item: item[0], reverse=True)
    ordered_jobs = [job for _, job in reranked]
    ordered_jobs.extend(job for index, job in enumerate(candidates) if index not in scored_indexes)
    return ordered_jobs + jobs[10:]


async def rerank_jobs_with_gemini(profile: str, jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Backward-compatible wrapper for callers using the original reranker name."""
    return await rerank_jobs(profile, jobs)


@app.get("/api/jobs/search")
async def search_saved_jobs(
    q: str = "",
    limit: int = 50,
    profile: str = "",
    rerank: bool = False,
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Retrieve a broad hybrid pool and optionally rerank its top candidates."""
    jobs = [
        job for job in import_jobs_from_csv("job_matches.csv", current_user["uid"])
        if job.get("User Dismissed", "").lower() != "yes"
    ]
    bounded_limit = max(1, min(limit, 100))
    cache_key = (current_user["uid"], q.strip(), profile.strip(), rerank)
    cache_ttl = max(60, int(os.getenv("SEARCH_RESULTS_CACHE_TTL_SECONDS", "21600")))
    cached = SEARCH_RESULT_CACHE.get(cache_key)
    if cached and time.time() - cached[0] < cache_ttl:
        return JSONResponse({"jobs": cached[1][:bounded_limit], "total": len(jobs), "cached": True})
    retrieval_limit = min(max(bounded_limit, 50), 100)
    retrieved_jobs = hybrid_search_jobs(jobs, q, retrieval_limit)
    if rerank and profile.strip() and retrieved_jobs:
        retrieved_jobs = await rerank_jobs(profile, retrieved_jobs)
    SEARCH_RESULT_CACHE[cache_key] = (time.time(), retrieved_jobs)
    return JSONResponse({"jobs": retrieved_jobs[:bounded_limit], "total": len(jobs)})


@app.post("/api/jobs/feedback")
async def submit_match_feedback(
    job_title: str = Form(...),
    company: str = Form(""),
    feedback_type: str = Form(...),
    notes: str = Form(""),
    fit_score: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    allowed_feedback = {"good_match", "poor_match", "irrelevant", "job_expired", "invalid_listing", "too_senior", "wrong_tech_stack", "wrong_location"}
    if feedback_type not in allowed_feedback:
        raise HTTPException(status_code=400, detail="Unsupported feedback type")
    record = firebase_utils.save_match_feedback(
        current_user["uid"], job_title, company, feedback_type, notes, fit_score
    )
    if feedback_type in {"job_expired", "invalid_listing"}:
        update_job_status_in_csv(
            job_title,
            company,
            "Expired",
            "job_matches.csv",
            current_user["uid"],
        )
    return JSONResponse({"feedback": record})


@app.get("/api/preferences/roles")
async def get_role_preferences(current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    return JSONResponse(firebase_utils.fetch_role_preferences(current_user["uid"]))


@app.put("/api/preferences/roles")
async def put_role_preferences(
    target_roles: str = Form(""),
    excluded_roles: str = Form(""),
    target_location: str = Form(""),
    max_posting_age_days: int = Form(7),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    split_roles = lambda value: list(dict.fromkeys(role.strip() for role in re.split(r"[,\n]", value) if role.strip()))
    record = firebase_utils.save_role_preferences(current_user["uid"], split_roles(target_roles), split_roles(excluded_roles), target_location, max_posting_age_days)
    return JSONResponse(record)


@app.post("/api/jobs/import-csv")
async def import_existing_csv_jobs(current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    """Migrate local job_matches.csv records into the signed-in user's store."""
    local_jobs = import_local_jobs_from_csv("job_matches.csv")
    if not local_jobs:
        return JSONResponse({"imported": 0, "message": "No local job_matches.csv records found."})
    append_jobs_to_csv(local_jobs, "job_matches.csv", user_id=current_user["uid"])
    return JSONResponse({"imported": len(local_jobs), "message": f"Imported {len(local_jobs)} CSV jobs."})


@app.post("/api/jobs/dismiss")
async def dismiss_job(
    job_title: str = Form(...),
    company: str = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Hide a job from the listing while retaining it in CSV for deduplication."""
    dismissed = dismiss_job_in_csv(job_title, company, "job_matches.csv", current_user["uid"])
    if not dismissed:
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"dismissed": True, "message": "Job dismissed and retained in CSV"})


@app.post("/api/jobs/dismiss-expired")
async def dismiss_all_expired_jobs(
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Dismiss all jobs marked as Expired, keeping them for deduplication."""
    count = dismiss_expired_jobs_in_csv("job_matches.csv", current_user["uid"])
    return JSONResponse({"dismissed": count})


@app.post("/api/jobs/update-url")
async def update_job_url(
    job_title: str = Form(...),
    company: str = Form(...),
    url: str = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Save a user-confirmed preferred listing URL for a job."""
    if not re.match(r"^https?://[^\s]+$", url.strip(), re.IGNORECASE):
        raise HTTPException(status_code=400, detail="Enter a valid http(s) URL")
    updated = update_job_url_in_csv(job_title, company, url.strip(), "job_matches.csv", current_user["uid"])
    if not updated:
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"updated": True, "url": url.strip()})


@app.post("/api/jobs/applied")
async def update_job_applied(
    job_title: str = Form(...),
    company: str = Form(...),
    applied: bool = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Persist the user's applied/not-applied status for a job."""
    updated = update_job_applied_in_csv(job_title, company, applied, "job_matches.csv", current_user["uid"])
    if not updated:
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"updated": True, "applied": applied})


@app.post("/api/jobs/status")
async def update_job_status(
    job_title: str = Form(...),
    company: str = Form(...),
    status: str = Form(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
) -> JSONResponse:
    """Persist a manual verification status for a job row."""
    if not status.strip():
        raise HTTPException(status_code=400, detail="Status is required")

    updated = update_job_status_in_csv(job_title, company, status.strip(), "job_matches.csv", current_user["uid"])
    if not updated:
        raise HTTPException(status_code=404, detail="Job not found")
    return JSONResponse({"updated": True, "status": status.strip()})


@app.post("/api/jobs/verify/start")
async def start_job_verification(current_user: Dict[str, Any] = Depends(get_current_user)) -> JSONResponse:
    """Start AI verification of the saved job listings."""
    verification_id = f"verify-{int(time.time() * 1000)}-{uuid.uuid4().hex[:8]}"
    ACTIVE_SEARCHES[verification_id] = {
        "status": "queued",
        "progress": 0,
        "message": "Starting saved job verification...",
        "logs": ["Starting saved job verification..."],
        "user_id": current_user["uid"],
    }
    asyncio.create_task(run_saved_job_verification(verification_id, current_user["uid"]))
    return JSONResponse({"verification_id": verification_id, "status": "queued"})
